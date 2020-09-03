#!/usr/bin/python
# -*- coding:UTF-8 -*-
'''
Author: zhangpeiyu
Date: 2020-09-01 21:51:45
LastEditTime: 2020-09-03 21:56:18
Description: 我不是诗人，所以，只能够把爱你写进程序，当作不可解的密码，作为我一个人知道的秘密。
'''
import file_to_json

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    document = Document()
    document.add_heading(u'接口文档', 0)

    # 读取json数据
    json_data = file_to_json.getJson('api.json')

    root = '/root'
    tags = json_data['tags']
    paths = json_data['paths']
    definitions = json_data['definitions']

    temp = {}
    for url, url_info in paths.items():
        for method, req_info in url_info.items():
            for prop, val in req_info.items():
                if prop == 'tags':
                    if val[0] in temp:
                        arr = temp[val[0]]
                        item2 = {}
                        item2[method] = req_info
                        item = {}
                        item[url] = item2
                        arr.append(item)
                        temp[val[0]] = arr
                        break
                    else:
                        arr = []
                        item2 = {}
                        item2[method] = req_info
                        item = {}
                        item[url] = item2
                        arr.append(item)
                        temp[val[0]] = arr
                        break
    num = 0
    for index1, item1 in enumerate(tags):
        num = index1
        first_title = item1['name']
        # 一级标题
        document.add_heading(''.join([str(index1 + 1), '.', first_title]), 1)

        # 二级标题
        second_title(document, temp, definitions, first_title, index1, root)

    # 实体
    definition(document, definitions, num+1)

    # 保存文档
    document.save('test.docx')


def definition(document, definitions, index1):
    document.add_heading(''.join([str(index1 + 1), '.', 'definitions']), 1)
    num = 0
    for beanName, beanInfo in definitions.items():
        document.add_heading(''.join([str(num + 1), '.', beanName]), 2)
        table = document.add_table(
            rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'字段名'
        hdr_cells[1].text = u'参数位置'
        hdr_cells[2].text = u'字段类型'
        hdr_cells[3].text = u'是否必填'
        hdr_cells[4].text = u'备注'
        prop_to_row(table, beanInfo['properties'], '')


def second_title(document, temp, definitions, first_title, index1, root):
    for tag, values in temp.items():
        if tag == first_title:
            for index2, item2 in enumerate(values):
                ind = index2
                for url, url_info in item2.items():
                    for method, method_info in url_info.items():
                        for prop, val in method_info.items():
                            if 'summary' == prop:
                                document.add_heading(
                                    ''.join([str(index1 + 1), '.', str(ind + 1), '.', val]), 2)
                                break
                        # 添加段落
                        paragraph = document.add_paragraph(u'请求方式：')
                        paragraph.add_run(method)
                        paragraph.add_run('\n')

                        paragraph.add_run(u'请求链接：')
                        paragraph.add_run(root + url)
                        paragraph.add_run('\n')

                        paragraph.add_run(u'功能描述：')
                        paragraph.add_run(method_info['description'])
                        paragraph.add_run('\n')

                        paragraph.add_run(u'请求类型（Content-Type）：')
                        paragraph.add_run(
                            '、'.join(method_info.get('consumes', [])))
                        paragraph.add_run('\n')

                        paragraph.add_run(u'响应类型（Content-Type）：')
                        paragraph.add_run(
                            '、'.join(method_info.get('produces', [])))
                        paragraph.add_run('\n')
                        parameters = method_info.get('parameters', [])
                        if len(parameters) > 0:
                            for row_index, row_data in enumerate(parameters):
                                if 'body' != row_data.get('in', ''):
                                    if row_index == 0:
                                        paragraph.add_run(u'请求参数：')
                                        table = document.add_table(
                                            rows=1, cols=5)
                                        hdr_cells = table.rows[0].cells
                                        hdr_cells[0].text = u'字段名'
                                        hdr_cells[1].text = u'参数位置'
                                        hdr_cells[2].text = u'字段类型'
                                        hdr_cells[3].text = u'是否必填'
                                        hdr_cells[4].text = u'备注'
                                    table.add_row()
                                    hdr_cells = table.rows[row_index + 1].cells
                                    hdr_cells[0].text = row_data.get(
                                        'name', '')
                                    hdr_cells[1].text = row_data.get(
                                        'in', '')
                                    if row_data.get(
                                            'type', '').capitalize() == 'Integer':
                                        if row_data.get('type', '') == 'int32':
                                            hdr_cells[2].text = row_data.get(
                                                'type', '').capitalize()
                                        else:
                                            hdr_cells[2].text = 'Long'
                                    else:
                                        hdr_cells[2].text = row_data.get(
                                            'type', '').capitalize()
                                    hdr_cells[3].text = str(
                                        row_data.get('required', False))
                                    hdr_cells[4].text = row_data.get(
                                        'description', '')
                                else:
                                    paragraph.add_run(u'请求参数：')
                                    table = document.add_table(
                                        rows=1, cols=5)
                                    hdr_cells = table.rows[0].cells
                                    hdr_cells[0].text = u'字段名'
                                    hdr_cells[1].text = u'参数位置'
                                    hdr_cells[2].text = u'字段类型'
                                    hdr_cells[3].text = u'是否必填'
                                    hdr_cells[4].text = u'备注'
                                    properties = {}
                                    if '$ref' in row_data['schema']:
                                        ref = row_data['schema']['$ref'].replace(
                                            '#/definitions/', '')
                                        properties = definitions[ref]['properties']
                                    else:
                                        ref = row_data['schema']['items']['$ref'].replace(
                                            '#/definitions/', '')
                                        properties = definitions[ref]['properties']

                                    prop_to_row(table, properties, 'body')

                        else:
                            paragraph.add_run(u'请求参数：')
                            paragraph.add_run(u'无')

                        paragraph = document.add_paragraph(u'响应参数：')
                responses(definitions, method_info['responses'],
                          document, paragraph)

# 响应参数


def responses(definitions, responses, document, paragraph):
    for status, res_info in responses.items():
        if status == '200':
            table = document.add_table(rows=1, cols=5)
            if '$ref' in res_info['schema']:
                ref = res_info['schema']['$ref'].replace(
                    '#/definitions/', '').capitalize()
                for key, prop in definitions.items():
                    if ref.lower() == key.lower():
                        properties = prop['properties']
                        table = document.add_table(
                            rows=1, cols=5)
                        prop_to_row(table, properties, '')
                    print(200)
            elif 'type' in res_info['schema']:
                print(202, res_info['schema'])
                if 'additionalProperties' in res_info['schema']:
                    prop = {}
                    prop['additionalProperties'] = res_info['schema']['additionalProperties']
                    prop_to_row(table, prop, '')
                elif 'items' in res_info['schema']:
                    print(208, res_info['schema'])
                    ref = res_info['schema']['items']['$ref'].replace(
                        '#/definitions/', '').capitalize()
                    for key, prop in definitions.items():
                        if ref.lower() == key.lower():
                            properties = prop['properties']
                            table = document.add_table(
                                rows=1, cols=5)
                            prop_to_row(table, properties, '')
                        else:
                            print('217')
                else:
                    prop = {}
                    prop['response'] = res_info['schema']
                    prop_to_row(table, prop, '')
            else:
                print('response')
                pass

        else:
            paragraph.add_run(u'无')


def prop_to_row(table, properties, pos):
    for property_index, property_value in enumerate(properties):
        table.add_row()
        hdr_cells = table.rows[property_index + 1].cells
        if 'type' in properties[property_value]:
            hdr_cells[0].text = property_value
            hdr_cells[1].text = pos
            hdr_cells[2].text = properties[property_value].get(
                'type', '').capitalize()
            hdr_cells[3].text = str(
                properties[property_value].get('required', False))
            hdr_cells[4].text = properties[property_value].get(
                'description', '')
        elif '$ref' in properties[property_value]:
            ref = properties[property_value].get(
                '$ref', '').replace('#/definitions/', '').capitalize()
            hdr_cells[0].text = property_value
            hdr_cells[1].text = pos
            hdr_cells[2].text = ref.capitalize()
            hdr_cells[3].text = str(
                properties[property_value].get('required', False))
            hdr_cells[4].text = properties[property_value].get(
                'description', '')
        elif True:
            print('pass')
            pass
            # body用string接收

        else:
            hdr_cells[0].text = property_value
            hdr_cells[1].text = pos
            hdr_cells[2].text = properties[property_value].get(
                '$ref', '').replace('#/definitions/', '').capitalize()
            hdr_cells[3].text = str(
                properties[property_value].get('required', False))
            hdr_cells[4].text = properties[property_value].get(
                'description', '')


if __name__ == '__main__':
    main()
