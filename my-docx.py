from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

#打开文档
document = Document()

#加入不同等级的标题
x = document.add_heading('python-docx文档',0)
print(x)
#print(help(document))
run = document.add_heading(u'一',1).add_run(u"添加文档标题")

run.font.name=u'微软雅黑'
#paragraph_format = run.run_format
#paragraph_format.line_spacing = 4 # 1.75倍行间距
document.add_heading(u'二级标题',2)
#添加文本
paragraph = document.add_paragraph(u'添加了文本')
print(paragraph)
#设置字号
run = paragraph.add_run(u'设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加表格
table = document.add_table(rows=3,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text="第一列"
hdr_cells[1].text="第二列"
hdr_cells[2].text="第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = 'aerszvfdgx'
hdr_cells[2].text = 'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '3'
hdr_cells[1].text = 'cafdwvaef'
hdr_cells[2].text = 'aabs zfgf'
document.add_heading(text=u'这是文档标题', level=0)

document.add_heading(text=u'这是一级标题', level=1)

document.add_heading(text=u'这是二级标题', level=2)

document.add_heading(text=u'这是三级标题', level=3)

document.add_heading(text=u'这是四级标题', level=4)

document.add_heading(text=u'这是五级标题', level=5)

document.add_heading(text=u'这是六级标题', level=6)

document.add_heading(text=u'这是七级标题', level=7)

document.add_heading(text=u'这是八级标题', level=8)

document.add_heading(text=u'这是九级标题', level=9)

#添加文本
paragraph = document.add_paragraph(u'添加了文本')
#设置字号
run = paragraph.add_run(u'设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'

#设置中文字体
run = paragraph.add_run(u'设置中文字体，')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    u'有序列表元素1',style='List Number'
)
document.add_paragraph(
    u'有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    u'无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
document.add_picture('jdb.jpg',width=Inches(1.25))

#增加表格
table = document.add_table(rows=3,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text="第一列"
hdr_cells[1].text="第二列"
hdr_cells[2].text="第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = 'aerszvfdgx'
hdr_cells[2].text = 'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '3'
hdr_cells[1].text = 'cafdwvaef'
hdr_cells[2].text = 'aabs zfgf'

#增加分页
document.add_page_break()
document.add_page_break()
document.add_page_break()


#保存文件
document.save('demo.docx')
